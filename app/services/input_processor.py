"""Input Processor - Multiple formats se requirements extract karo"""

import os
import json
import logging
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
import asyncio

# Document parsers
import PyPDF2
from docx import Document
import openpyxl
import pandas as pd
import markdown
from bs4 import BeautifulSoup

# Audio/Video processing
import speech_recognition as sr
from pydub import AudioSegment
import moviepy.editor as mp
from app.core.llm import get_llm, invoke_llm_with_tracing

# LLM for summarization
from langchain_openai import ChatOpenAI
from langchain.prompts import PromptTemplate

logger = logging.getLogger(__name__)


class InputProcessor:
    """
    Multiple format inputs ko process karta hai:
    - PDF, Word, Excel, Text files
    - Audio/Video recordings
    - URLs (Google Docs, Confluence, etc.)
    - Direct text input
    """
    
    def __init__(self):
        self.llm = get_llm()
    
    async def process_inputs(self, inputs: List[Dict[str, Any]]) -> Dict[str, Any]:
        """
        Main entry point - multiple inputs process karo
        
        inputs = [
            {"type": "text", "content": "Build an AI chatbot..."},
            {"type": "file", "path": "meeting_notes.pdf"},
            {"type": "audio", "path": "client_call.mp3"},
            {"type": "url", "url": "https://docs.google.com/..."}
        ]
        """
        
        all_extracted = []
        
        for input_item in inputs:
            extracted = await self._process_single_input(input_item)
            if extracted:
                all_extracted.append(extracted)
        
        # Consolidate all inputs into unified requirements
        consolidated = await self._consolidate_requirements(all_extracted)
        
        return {
            "raw_inputs": all_extracted,
            "consolidated_idea": consolidated["idea"],
            "requirements": consolidated["requirements"],
            "constraints": consolidated["constraints"],
            "stakeholders": consolidated["stakeholders"],
            "source_count": len(all_extracted)
        }
    
    async def _process_single_input(self, input_item: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        """Single input process karo based on type"""
        
        input_type = input_item.get("type", "text")
        
        try:
            if input_type == "text":
                return await self._process_text(input_item.get("content", ""))
            
            elif input_type == "file":
                file_path = input_item.get("path")
                ext = Path(file_path).suffix.lower()
                
                if ext == ".pdf":
                    return await self._process_pdf(file_path)
                elif ext in [".docx", ".doc"]:
                    return await self._process_word(file_path)
                elif ext in [".xlsx", ".xls"]:
                    return await self._process_excel(file_path)
                elif ext == ".txt":
                    return await self._process_text_file(file_path)
                elif ext == ".md":
                    return await self._process_markdown(file_path)
                elif ext in [".mp3", ".wav", ".m4a"]:
                    return await self._process_audio(file_path)
                elif ext in [".mp4", ".mov", ".avi"]:
                    return await self._process_video(file_path)
                else:
                    logger.warning(f"Unsupported file type: {ext}")
                    return None
            
            elif input_type == "url":
                return await self._process_url(input_item.get("url"))
            
            elif input_type == "audio":
                return await self._process_audio(input_item.get("path"))
            
            else:
                logger.warning(f"Unknown input type: {input_type}")
                return None
                
        except Exception as e:
            logger.error(f"Failed to process {input_type}: {e}")
            return {"error": str(e), "type": input_type}
    
    async def _process_text(self, content: str) -> Dict[str, Any]:
        """Simple text input process"""
        return {
            "type": "text",
            "content": content,
            "summary": await self._summarize(content),
            "source": "direct_text"
        }
    
    async def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """PDF file se text extract karo"""
        text = ""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() or ""
        
        return {
            "type": "pdf",
            "content": text,
            "summary": await self._summarize(text),
            "source": file_path,
            "page_count": len(reader.pages)
        }
    
    async def _process_word(self, file_path: str) -> Dict[str, Any]:
        """Word document se text extract karo"""
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Extract tables also
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        return {
            "type": "word",
            "content": text,
            "tables": tables,
            "summary": await self._summarize(text),
            "source": file_path
        }
    
    async def _process_excel(self, file_path: str) -> Dict[str, Any]:
        """Excel file se data extract karo"""
        df = pd.read_excel(file_path, sheet_name=None)
        
        all_text = []
        sheets_data = {}
        
        for sheet_name, data in df.items():
            sheets_data[sheet_name] = data.to_dict(orient='records')
            all_text.append(data.to_string())
        
        combined_text = "\n".join(all_text)
        
        return {
            "type": "excel",
            "content": combined_text,
            "sheets": sheets_data,
            "summary": await self._summarize(combined_text),
            "source": file_path
        }
    
    async def _process_text_file(self, file_path: str) -> Dict[str, Any]:
        """Plain text file read karo"""
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        return {
            "type": "text_file",
            "content": text,
            "summary": await self._summarize(text),
            "source": file_path
        }
    
    async def _process_markdown(self, file_path: str) -> Dict[str, Any]:
        """Markdown file se text extract karo"""
        with open(file_path, 'r', encoding='utf-8') as file:
            md_text = file.read()
        
        # Convert to HTML then to text
        html = markdown.markdown(md_text)
        soup = BeautifulSoup(html, 'html.parser')
        text = soup.get_text()
        
        return {
            "type": "markdown",
            "content": text,
            "markdown_raw": md_text,
            "summary": await self._summarize(text),
            "source": file_path
        }
    
    async def _process_audio(self, file_path: str) -> Dict[str, Any]:
        """Audio file se speech to text"""
        try:
            # Convert to WAV if needed
            audio = AudioSegment.from_file(file_path)
            wav_path = file_path + ".wav"
            audio.export(wav_path, format="wav")
            
            # Speech recognition
            recognizer = sr.Recognizer()
            with sr.AudioFile(wav_path) as source:
                audio_data = recognizer.record(source)
                text = recognizer.recognize_google(audio_data)
            
            # Cleanup
            os.remove(wav_path)
            
            return {
                "type": "audio",
                "content": text,
                "summary": await self._summarize(text),
                "source": file_path,
                "duration": len(audio) / 1000  # seconds
            }
            
        except Exception as e:
            logger.error(f"Audio processing failed: {e}")
            return {
                "type": "audio",
                "error": str(e),
                "source": file_path
            }
    
    async def _process_video(self, file_path: str) -> Dict[str, Any]:
        """Video file se audio extract karo phir transcribe"""
        try:
            # Extract audio from video
            video = mp.VideoFileClip(file_path)
            audio_path = file_path + ".wav"
            video.audio.write_audiofile(audio_path)
            
            # Transcribe audio
            recognizer = sr.Recognizer()
            with sr.AudioFile(audio_path) as source:
                audio_data = recognizer.record(source)
                text = recognizer.recognize_google(audio_data)
            
            # Cleanup
            os.remove(audio_path)
            video.close()
            
            return {
                "type": "video",
                "content": text,
                "summary": await self._summarize(text),
                "source": file_path,
                "duration": video.duration
            }
            
        except Exception as e:
            logger.error(f"Video processing failed: {e}")
            return {
                "type": "video",
                "error": str(e),
                "source": file_path
            }
    
    async def _process_url(self, url: str) -> Dict[str, Any]:
        """URL se content fetch karo"""
        try:
            import aiohttp
            from urllib.parse import urlparse
            
            # Check if it's Google Docs
            if "docs.google.com" in url:
                return await self._process_google_docs(url)
            
            # Generic web page
            async with aiohttp.ClientSession() as session:
                async with session.get(url) as response:
                    html = await response.text()
                    soup = BeautifulSoup(html, 'html.parser')
                    
                    # Remove script and style elements
                    for script in soup(["script", "style"]):
                        script.decompose()
                    
                    text = soup.get_text()
                    lines = (line.strip() for line in text.splitlines())
                    chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                    text = ' '.join(chunk for chunk in chunks if chunk)
            
            return {
                "type": "url",
                "content": text[:50000],  # Limit length
                "summary": await self._summarize(text[:10000]),
                "source": url
            }
            
        except Exception as e:
            logger.error(f"URL processing failed: {e}")
            return {"type": "url", "error": str(e), "source": url}
    
    async def _process_google_docs(self, url: str) -> Dict[str, Any]:
        """Google Docs specific handling"""
        # Note: This requires Google Docs API
        # For now, return placeholder
        return {
            "type": "google_docs",
            "content": "Google Docs content - API integration needed",
            "summary": "Google Docs document",
            "source": url,
            "note": "Full integration requires Google Docs API"
        }
    
    async def _summarize(self, text: str, max_length: int = 500) -> str:
        """Long text ko summarize karo for quick understanding"""
        if not text or len(text) < 200:
            return text
        
        prompt = PromptTemplate(
            input_variables=["text"],
            template="""
            Summarize the following text in 2-3 sentences. Focus on key requirements, goals, and constraints:
            
            Text: {text}
            
            Summary:
            """
        )
        
        chain = prompt | self.llm
        result = await chain.ainvoke({"text": text[:3000]})  # Limit input
        return result.content
    
    async def _consolidate_requirements(self, all_inputs: List[Dict]) -> Dict[str, Any]:
        """Multiple inputs se consolidated requirements extract karo"""
        
        # Combine all content
        combined_content = "\n\n---\n\n".join([
            f"Source: {inp.get('source', inp.get('type'))}\n{inp.get('content', '')}"
            for inp in all_inputs if inp.get('content')
        ])
        
        prompt = PromptTemplate(
            input_variables=["content"],
            template="""
            You are a Product Manager consolidating requirements from multiple sources.
            
            Here are all the inputs (meetings, documents, emails, etc.):
            
            {content}
            
            Extract and consolidate into:
            
            1. **Idea**: One clear sentence describing what to build
            2. **Requirements**: List of functional requirements (be specific)
            3. **Constraints**: Technical, budget, timeline, regulatory constraints
            4. **Stakeholders**: Who are the key people involved?
            5. **Open Questions**: What information is missing?
            
            Return ONLY valid JSON:
            {{
                "idea": "string",
                "requirements": ["req1", "req2"],
                "constraints": ["constraint1", "constraint2"],
                "stakeholders": ["person1", "person2"],
                "open_questions": ["question1", "question2"]
            }}
            """
        )
        
        chain = prompt | self.llm
        result = await chain.ainvoke({"content": combined_content[:10000]})
        
        try:
            import json
            # Clean response if needed
            response_text = result.content
            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0]
            elif "```" in response_text:
                response_text = response_text.split("```")[1].split("```")[0]
            
            consolidated = json.loads(response_text.strip())
            return consolidated
            
        except Exception as e:
            logger.error(f"Failed to parse consolidated requirements: {e}")
            return {
                "idea": combined_content[:500],
                "requirements": [],
                "constraints": [],
                "stakeholders": [],
                "open_questions": []
            }